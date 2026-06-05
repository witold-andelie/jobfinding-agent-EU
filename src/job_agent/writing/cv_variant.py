"""Tailored CV variant — re-emphasise (never invent) for a specific job, export .docx.

Same hard rule as the cover letter: nothing fabricated. The variant only (a) reorders
the candidate's real skills so the ones the job wants come first, and (b) rewrites the
*summary* to target the role. Experience and education are carried over verbatim from
the parsed CV. Rendering uses python-docx (the `docx` extra).
"""

from __future__ import annotations

import re
import unicodedata
from pathlib import Path

from job_agent.models.job import Job
from job_agent.parsing.cv import ParsedCV

_WORD = re.compile(r"[a-zA-ZÀ-ſ]{3,}")

_SUMMARY_PROMPT = (
    "Write a 2-3 sentence professional CV summary for this candidate targeting the "
    "specific job below. Ground it ONLY in their real background; do NOT invent "
    "anything. Be concrete and concise.\n"
    "Candidate field: {field}; skills: {skills}; experience: {years} years.\n"
    "Target job: {title} at {company}.\nJob description:\n{description}\n\n"
    "Return only the summary text."
)


def _tokens(text: str) -> set[str]:
    return {t.lower() for t in _WORD.findall(text or "")}


def rank_skills(skills: list[str], job: Job, matched: list[str] | None = None) -> list[str]:
    """Reorder skills: gap-analysis matches first, then those overlapping the job text.

    Deterministic and stable; never adds or drops a skill (no fabrication).
    """
    matched_lower = {m.lower() for m in (matched or [])}
    job_tokens = _tokens(f"{job.title} {job.description}")

    def rank(skill: str) -> tuple[int, int]:
        if skill.lower() in matched_lower:
            return (2, 0)
        overlap = len(_tokens(skill) & job_tokens)
        return (1, overlap) if overlap else (0, 0)

    return sorted(skills, key=rank, reverse=True)


def tailored_summary(parsed: ParsedCV, job: Job, ask) -> str:  # ask: Callable[[str], str]
    """DeepSeek summary targeting the role, grounded in the real profile."""
    prompt = _SUMMARY_PROMPT.format(
        field=parsed.field or "(unspecified)",
        skills=", ".join(parsed.skills) or "(none listed)",
        years=parsed.years_experience,
        title=job.title,
        company=job.company,
        description=(job.description or "")[:3000],
    )
    return ask(prompt).strip()


def build_cv_docx(parsed: ParsedCV, ranked_skills: list[str], summary: str, path: str | Path) -> Path:
    """Render a tailored CV to ``path`` (.docx). Experience/education kept verbatim."""
    from docx import Document  # lazy import — only needed when exporting

    doc = Document()
    doc.add_heading(parsed.name or "Curriculum Vitae", level=0)
    if parsed.contact:
        doc.add_paragraph(parsed.contact)

    doc.add_heading("Profile", level=1)
    doc.add_paragraph(summary)

    doc.add_heading("Key Skills", level=1)
    for skill in ranked_skills:
        doc.add_paragraph(skill, style="List Bullet")

    if parsed.languages:
        doc.add_heading("Languages", level=1)
        doc.add_paragraph(", ".join(parsed.languages))

    for title, lines in (("Experience", parsed.experience), ("Education", parsed.education)):
        if lines:
            doc.add_heading(title, level=1)
            for line in lines:
                doc.add_paragraph(line, style="List Bullet")

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def _slug(text: str) -> str:
    ascii_text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode()
    return re.sub(r"[^a-z0-9]+", "-", ascii_text.lower()).strip("-") or "job"


def generate_cv_variant(
    parsed: ParsedCV,
    job: Job,
    ask,  # ask: Callable[[str], str]
    *,
    artifacts_dir: str | Path = "artifacts",
    matched: list[str] | None = None,
) -> Path:
    """Full variant: re-rank skills + tailored summary + .docx export. Returns the path."""
    ranked = rank_skills(parsed.skills, job, matched)
    summary = tailored_summary(parsed, job, ask)
    path = Path(artifacts_dir) / f"CV_{_slug(job.company)}_{job.external_id}.docx"
    return build_cv_docx(parsed, ranked, summary, path)
